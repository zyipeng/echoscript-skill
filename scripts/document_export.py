#!/usr/bin/env python3
"""Export canonical EchoScript Markdown to DOCX and PDF, then validate it."""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
from pathlib import Path
import re
import subprocess
import sys
from typing import Any


class ExportError(RuntimeError):
    pass


DOCUMENT_FONT = "Hiragino Sans GB"


def has_document_packages() -> bool:
    return bool(importlib.util.find_spec("docx") and importlib.util.find_spec("pypdf") and importlib.util.find_spec("reportlab"))


def find_document_python() -> Path | None:
    override = os.environ.get("ECHOSCRIPT_DOCUMENT_PYTHON")
    candidates: list[Path] = []
    if override:
        candidates.append(Path(override).expanduser())
    runtime_root = Path.home() / ".cache" / "codex-runtimes"
    candidates.extend(sorted(runtime_root.glob("*/dependencies/python/bin/python3"), reverse=True))
    for candidate in candidates:
        if not candidate.is_file() or candidate.resolve() == Path(sys.executable).resolve():
            continue
        result = subprocess.run(
            [str(candidate), "-c", "import docx, pypdf, reportlab"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if result.returncode == 0:
            return candidate
    return None


def maybe_relaunch() -> int | None:
    if has_document_packages() or os.environ.get("ECHOSCRIPT_DOCUMENT_RUNTIME") == "1":
        return None
    candidate = find_document_python()
    if not candidate:
        raise ExportError("缺少 python-docx/pypdf，也没有找到 Codex 文档运行时")
    environment = os.environ.copy()
    environment["ECHOSCRIPT_DOCUMENT_RUNTIME"] = "1"
    result = subprocess.run([str(candidate), str(Path(__file__).resolve()), *sys.argv[1:]], env=environment)
    return result.returncode


def safe_name(value: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "-", value)
    name = re.sub(r"\s+", " ", name).strip(" .-")
    return name[:100] or "EchoScript-document"


def strip_frontmatter(text: str) -> str:
    if not text.startswith("---\n"):
        return text
    end = text.find("\n---\n", 4)
    return text[end + 5 :] if end >= 0 else text


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def special_line(lines: list[str], index: int) -> bool:
    line = lines[index]
    stripped = line.strip()
    if not stripped:
        return True
    if re.match(r"^#{1,6}\s+", stripped) or re.match(r"^[-*+]\s+", stripped) or re.match(r"^\d+[.)]\s+", stripped):
        return True
    if stripped.startswith(">") or stripped.startswith("```") or re.fullmatch(r"[-*_]{3,}", stripped):
        return True
    return index + 1 < len(lines) and "|" in stripped and is_table_separator(lines[index + 1])


def parse_markdown(text: str) -> list[dict[str, Any]]:
    lines = strip_frontmatter(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            body: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index])
                index += 1
            index += 1
            blocks.append({"type": "code", "language": language, "text": "\n".join(body)})
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            blocks.append({"type": "heading", "level": len(heading.group(1)), "text": heading.group(2).strip()})
            index += 1
            continue
        if index + 1 < len(lines) and "|" in stripped and is_table_separator(lines[index + 1]):
            rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            blocks.append({"type": "bullet", "text": bullet.group(1).strip()})
            index += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            blocks.append({"type": "number", "text": numbered.group(1).strip()})
            index += 1
            continue
        if stripped.startswith(">"):
            body: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                body.append(lines[index].strip()[1:].lstrip())
                index += 1
            blocks.append({"type": "quote", "text": "\n".join(body)})
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            blocks.append({"type": "rule"})
            index += 1
            continue
        paragraph = [stripped]
        index += 1
        while index < len(lines) and not special_line(lines, index):
            paragraph.append(lines[index].strip())
            index += 1
        blocks.append({"type": "paragraph", "text": " ".join(item for item in paragraph if item)})
    return blocks


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph: Any, text: str) -> None:
    from docx.shared import Pt

    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            paragraph.add_run(f"{link.group(1)} ({link.group(2)})" if link else token)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def configure_styles(document: Any) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = DOCUMENT_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.22
    title = document.styles["Title"]
    title.font.size = Pt(24)
    title.font.bold = True
    for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("EchoScript  ·  ")
    run.font.size = Pt(8)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def force_run_font(run: Any, font_name: str = DOCUMENT_FONT) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = font_name
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ["ascii", "hAnsi", "eastAsia", "cs"]:
        fonts.set(qn(f"w:{attribute}"), font_name)


def force_document_fonts(document: Any) -> None:
    paragraphs = list(document.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            force_run_font(run)


def create_docx(markdown: str, destination: Path, title_override: str | None = None) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    blocks = parse_markdown(markdown)
    document = Document()
    configure_styles(document)
    document.core_properties.title = title_override or next((block["text"] for block in blocks if block["type"] == "heading"), "EchoScript")
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = int(block["level"])
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(paragraph, title_override or block["text"])
            else:
                paragraph = document.add_heading(level=min(level - 1, 3))
                add_inline(paragraph, block["text"])
        elif kind == "paragraph":
            add_inline(document.add_paragraph(), block["text"])
        elif kind == "bullet":
            add_inline(document.add_paragraph(style="List Bullet"), block["text"])
        elif kind == "number":
            add_inline(document.add_paragraph(style="List Number"), block["text"])
        elif kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            run = paragraph.add_run(block["text"])
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            run.font.name = "Menlo"
            run.font.size = Pt(8.5)
        elif kind == "table":
            rows = block["rows"]
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table_properties = table._tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "6")
                border.set(qn("w:color"), "94A3B8")
                borders.append(border)
            table_properties.append(borders)
            for row_index, row in enumerate(rows):
                for column_index in range(width):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.width = Inches(6.5 / width)
                    margins = OxmlElement("w:tcMar")
                    for edge in ["top", "left", "bottom", "right"]:
                        margin = OxmlElement(f"w:{edge}")
                        margin.set(qn("w:w"), "100")
                        margin.set(qn("w:type"), "dxa")
                        margins.append(margin)
                    cell._tc.get_or_add_tcPr().append(margins)
                    cell.text = row[column_index] if column_index < len(row) else ""
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "E9EEF5")
                        cell._tc.get_or_add_tcPr().append(shading)
        elif kind == "rule":
            paragraph = document.add_paragraph()
            paragraph.add_run("────────────────────────").font.color.rgb = RGBColor(180, 180, 180)
    add_page_number(document.sections[0].footer.paragraphs[0])
    force_document_fonts(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def cjk_font_path() -> Path:
    candidates = [
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    ]
    candidates.extend(sorted(Path("/System/Library/AssetsV2/com_apple_MobileAsset_Font7").glob("*/AssetData/STHEITI.ttf")))
    candidates.extend([
        Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
        Path("/System/Library/Fonts/STHeiti Light.ttc"),
    ])
    for path in candidates:
        if path.is_file():
            return path
    raise ExportError("没有找到可嵌入 PDF 的中文字体")


def reportlab_inline(text: str) -> str:
    import html

    output: list[str] = []
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        output.append(html.escape(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            output.append(f"<b>{html.escape(token[2:-2])}</b>")
        elif token.startswith("`"):
            output.append(f"<font color='#334155'>{html.escape(token[1:-1])}</font>")
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                output.append(f"<link href='{html.escape(link.group(2), quote=True)}' color='#1D4ED8'>{html.escape(link.group(1))}</link>")
            else:
                output.append(html.escape(token))
        cursor = match.end()
    output.append(html.escape(text[cursor:]))
    return "".join(output).replace("\n", "<br/>")


def create_pdf(markdown: str, destination: Path, title_override: str | None = None) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "EchoScriptCJK"
    font_path = cjk_font_path()
    try:
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    except Exception as error:
        raise ExportError(f"无法加载 PDF 中文字体：{font_path}") from error
    pdfmetrics.registerFontFamily(font_name, normal=font_name, bold=font_name, italic=font_name, boldItalic=font_name)
    base = getSampleStyleSheet()
    normal = ParagraphStyle(
        "EchoNormal", parent=base["Normal"], fontName=font_name, fontSize=10.2,
        leading=15.2, textColor=colors.HexColor("#172033"), spaceAfter=6, wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "EchoTitle", parent=normal, fontSize=23, leading=30, textColor=colors.HexColor("#183B67"),
        spaceAfter=12, keepWithNext=True,
    )
    heading_styles = {
        2: ParagraphStyle("EchoH2", parent=normal, fontSize=16, leading=22, textColor=colors.HexColor("#183B67"), spaceBefore=12, spaceAfter=6, keepWithNext=True),
        3: ParagraphStyle("EchoH3", parent=normal, fontSize=12.5, leading=18, textColor=colors.HexColor("#244C78"), spaceBefore=9, spaceAfter=4, keepWithNext=True),
    }
    quote_style = ParagraphStyle("EchoQuote", parent=normal, leftIndent=8 * mm, textColor=colors.HexColor("#475569"), borderColor=colors.HexColor("#94A3B8"), borderWidth=0, borderPadding=3)
    code_style = ParagraphStyle("EchoCode", parent=normal, fontSize=8.7, leading=12, leftIndent=4 * mm, backColor=colors.HexColor("#F1F5F9"), borderPadding=5)
    blocks = parse_markdown(markdown)
    story: list[Any] = []
    number_index = 0
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = int(block["level"])
            text = title_override or block["text"] if level == 1 else block["text"]
            story.append(Paragraph(reportlab_inline(text), title_style if level == 1 else heading_styles.get(min(level, 3), heading_styles[3])))
        elif kind == "paragraph":
            story.append(Paragraph(reportlab_inline(block["text"]), normal))
        elif kind == "bullet":
            story.append(Paragraph(reportlab_inline(block["text"]), normal, bulletText="•"))
        elif kind == "number":
            number_index += 1
            story.append(Paragraph(reportlab_inline(block["text"]), normal, bulletText=f"{number_index}."))
        elif kind == "quote":
            story.append(Paragraph(reportlab_inline(block["text"]), quote_style))
        elif kind == "code":
            story.append(Paragraph(reportlab_inline(block["text"]), code_style))
        elif kind == "rule":
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#94A3B8"), spaceBefore=5, spaceAfter=7))
        elif kind == "table":
            rows = block["rows"]
            width = max(len(row) for row in rows)
            data = [[Paragraph(reportlab_inline(row[index] if index < len(row) else ""), normal) for index in range(width)] for row in rows]
            table = Table(data, colWidths=[(A4[0] - 36 * mm) / width] * width, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E9EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#64748B")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(KeepTogether([table, Spacer(1, 4)]))

    destination.parent.mkdir(parents=True, exist_ok=True)
    document_title = title_override or next((block["text"] for block in blocks if block["type"] == "heading"), "EchoScript")

    def footer(canvas: Any, document: Any) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawCentredString(A4[0] / 2, 10 * mm, f"EchoScript  ·  {document.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(destination), pagesize=A4, title=document_title, author="EchoScript",
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=17 * mm, bottomMargin=18 * mm,
    )
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)


def parse_formats(value: str) -> list[str]:
    formats = list(dict.fromkeys(item.strip().lower() for item in value.split(",") if item.strip()))
    invalid = [item for item in formats if item not in {"md", "docx", "pdf"}]
    if invalid:
        raise ExportError(f"不支持的格式：{', '.join(invalid)}")
    return formats


def export(args: argparse.Namespace) -> None:
    source = Path(args.markdown).expanduser().resolve()
    if not source.is_file():
        raise ExportError(f"Markdown 不存在：{source}")
    markdown = source.read_text(encoding="utf-8-sig")
    if len(markdown.strip()) < 20:
        raise ExportError("Markdown 内容过短，拒绝导出空文档")
    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    formats = parse_formats(args.formats)
    base = safe_name(args.name or source.stem)
    outputs: list[str] = []
    if "md" in formats:
        markdown_path = output_dir / f"{base}.md"
        markdown_path.write_text(markdown, encoding="utf-8")
        outputs.append(str(markdown_path))
    needs_docx = "docx" in formats
    docx_path = output_dir / f"{base}.docx"
    if needs_docx:
        create_docx(markdown, docx_path, args.title)
        if "docx" in formats:
            outputs.append(str(docx_path))
    if "pdf" in formats:
        pdf_path = output_dir / f"{base}.pdf"
        create_pdf(markdown, pdf_path, args.title)
        outputs.append(str(pdf_path))
    print(json.dumps({"ok": True, "outputs": outputs}, ensure_ascii=False, indent=2))


def validate_path(path: Path) -> dict[str, Any]:
    if not path.is_file() or path.stat().st_size == 0:
        raise ExportError(f"文件为空或不存在：{path}")
    suffix = path.suffix.lower()
    result: dict[str, Any] = {"path": str(path), "size": path.stat().st_size, "valid": True}
    if suffix == ".md":
        result["characters"] = len(path.read_text(encoding="utf-8-sig"))
    elif suffix == ".docx":
        from docx import Document
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
        if not text:
            raise ExportError(f"DOCX 没有可读正文：{path}")
        result.update({"paragraphs": len(document.paragraphs), "tables": len(document.tables), "characters": len(text)})
    elif suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        if not reader.pages:
            raise ExportError(f"PDF 没有页面：{path}")
        result["pages"] = len(reader.pages)
        result["extracted_characters"] = sum(len(page.extract_text() or "") for page in reader.pages)
    return result


def validate(args: argparse.Namespace) -> None:
    target = Path(args.target).expanduser().resolve()
    if target.is_dir():
        files = sorted(path for path in target.iterdir() if path.suffix.lower() in {".md", ".docx", ".pdf"})
    else:
        files = [target]
    if not files:
        raise ExportError(f"没有找到可验证文档：{target}")
    results = [validate_path(path) for path in files]
    print(json.dumps({"ok": True, "files": results}, ensure_ascii=False, indent=2))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="EchoScript Markdown document exporter")
    subparsers = parser.add_subparsers(dest="command", required=True)
    export_parser = subparsers.add_parser("export")
    export_parser.add_argument("markdown")
    export_parser.add_argument("--output-dir", required=True)
    export_parser.add_argument("--formats", default="md,docx,pdf")
    export_parser.add_argument("--name")
    export_parser.add_argument("--title")
    export_parser.set_defaults(handler=export)
    validate_parser = subparsers.add_parser("validate")
    validate_parser.add_argument("target")
    validate_parser.set_defaults(handler=validate)
    return parser


def main() -> int:
    try:
        relaunched = maybe_relaunch()
        if relaunched is not None:
            return relaunched
        args = build_parser().parse_args()
        args.handler(args)
        return 0
    except (ExportError, OSError, subprocess.SubprocessError, json.JSONDecodeError) as error:
        print(json.dumps({"ok": False, "error": str(error)}, ensure_ascii=False), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
